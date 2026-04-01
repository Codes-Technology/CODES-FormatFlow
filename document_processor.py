"""
DocumentProcessor — Signal-Based Classifier

"""

import os
import re
import time
import traceback
import tempfile
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from utils.style_manager import StyleManager
from utils.adobe_helper import adobe_pdf_extract
from utils.toc_manager import TocManager
from utils.cover_page_manager import CoverPageManager
from config import TEMPLATE_DOCX, ADOBE_CLIENT_ID, ADOBE_CLIENT_SECRET

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


class DocumentProcessor:
    def __init__(self, template_path: str, font_family: str = 'Calibri', font_size: int = 11, 
                 include_cover: bool = False, include_toc: bool = False):
        self.template_path = template_path
        self.style_manager = StyleManager(template_path, font_family, font_size)
        self.include_cover = include_cover
        self.include_toc = include_toc
        self.toc_manager = TocManager()
        self.cover_manager = CoverPageManager()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PUBLIC API
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def universal_extract(self, input_path: str, output_path: str) -> dict:
        """Entry point for uploaded files. Routes by extension."""
        try:
            ext = os.path.splitext(input_path)[1].lower()
            if ext == '.pdf':
                return self._pipeline_pdf(input_path, output_path)
            elif ext == '.docx':
                return self._pipeline_docx(input_path, output_path)
            return {'success': False, 'error': f'Unsupported file type: {ext}'}
        except Exception as e:
            traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def html_to_docx(self, html: str) -> Document:
        doc = Document(self.template_path)
        soup = BeautifulSoup(html, 'html.parser')

        # If there are NO block-level tags like p, h1, div, etc., 
        # then this is likely plain text from the 'Edit' box.
        # In this case, we use the Smart Classifier to detect headings.
        has_block_tags = bool(soup.find(['p', 'h1', 'h2', 'h3', 'h4', 'div', 'ul', 'ol', 'table']))

        if not has_block_tags:
            print("[DocumentProcessor] Detected plain text input (likely Edit), using Smart Classifier.")
            lines = html.splitlines()
            raw_lines = []
            for line in lines:
                t = line.strip()
                if t:
                    # Mock the signals expected by the classifier
                    raw_lines.append({
                        'text': t,
                        'style': 'Normal',
                        'is_bold': False,      # Plain text doesn't have bold signals
                        'run_size': 0,
                        'num_id': 0,
                        'num_lvl': 0
                    })
            self._h4_counters = {}
            self._build_from_signals(raw_lines, doc)
        else:
            print(f"[DocumentProcessor] Detected HTML input (length: {len(html)}), using tag-based parser.")
            processed_elements = set()

            def process_node(node, container):
                if node in processed_elements: return
                if not getattr(node, 'name', None):
                    clean_text = str(node).strip()
                    if clean_text:
                        if isinstance(container, Paragraph): container.add_run(clean_text)
                        else: container.add_paragraph(clean_text)
                    return

                if node.name in ('h1', 'h2', 'h3', 'h4'):
                    level = int(node.name[1])
                    para = container.add_heading('', level=level)
                    for child in node.contents: process_node(child, para)
                    processed_elements.add(node)
                elif node.name == 'p':
                    para = container.add_paragraph()
                    for child in node.contents: process_node(child, para)
                    processed_elements.add(node)
                elif node.name == 'table':
                    self._add_html_table(container, node)
                    processed_elements.add(node)
                elif node.name in ('ul', 'ol'):
                    style = 'List Bullet' if node.name == 'ul' else 'List Number'
                    for li in node.find_all('li', recursive=False):
                        para = container.add_paragraph(style=style)
                        for child in li.contents: process_node(child, para)
                    processed_elements.add(node)
                elif node.name in ('b', 'strong', 'i', 'em', 'u', 'span'):
                    if isinstance(container, Paragraph):
                        run = container.add_run()
                        if node.name in ('b', 'strong'): run.bold = True
                        if node.name in ('i', 'em'): run.italic = True
                        if node.name == 'u': run.underline = True
                        for child in node.contents:
                            if not getattr(child, 'name', None): run.text += str(child)
                            else: process_node(child, container)
                    else:
                        para = container.add_paragraph()
                        process_node(node, para)
                    processed_elements.add(node)
                elif node.name == 'div':
                    block_children = node.find(['p', 'h1', 'h2', 'h3', 'h4', 'ul', 'ol', 'table', 'div'])
                    if not block_children:
                        para = container.add_paragraph()
                        for child in node.contents: process_node(child, para)
                        processed_elements.add(node)
                    else:
                        for child in node.contents: process_node(child, container)
                else:
                    for child in node.contents: process_node(child, container)

            for child in soup.contents:
                process_node(child, doc)

        styled_doc = self.style_manager.apply_template_styles(doc)
        self._apply_final_features(styled_doc)
        return styled_doc

    def _apply_final_features(self, doc: Document):
        """Adds Cover Page and Table of Contents if requested."""
        if self.include_cover:
            self.cover_manager.create_cover_page(doc)

        if self.include_toc:
            self.toc_manager.insert_toc(doc)

    def _add_html_table(self, doc: Document, element):
        """
        Robust HTML table parser.
        Handles nested tables, merged cells, and complex HTML structures.
        """
        try:
            # Find the table element
            table = element
            
            # Create Word table with estimated dimensions
            rows = table.find_all('tr')
            num_rows = len(rows)
            
            # Estimate columns by finding max cells in any row
            num_cols = 0
            for row in rows:
                cells = row.find_all(['td', 'th'])
                num_cols = max(num_cols, len(cells))
            
            if num_rows == 0 or num_cols == 0:
                return
            
            # Create Word table
            word_table = doc.add_table(rows=num_rows, cols=num_cols)
            word_table.style = 'Table Grid'
            
            # Process each row
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                
                # Track merged cells
                col_offset = 0
                
                for cell in cells:
                    # Get cell text
                    cell_text = ''.join(cell.get_text(separator=' ', strip=True))
                    
                    # Get cell attributes
                    rowspan = int(cell.get('rowspan', 1))
                    colspan = int(cell.get('colspan', 1))
                    
                    # Find the target cell in Word table
                    target_row = i
                    target_col = col_offset
                    
                    # Adjust for previous merged cells
                    while target_row < num_rows and target_col < num_cols:
                        if word_table.cell(target_row, target_col).text.strip() == '':
                            break
                        target_col += 1
                    
                    if target_row >= num_rows or target_col >= num_cols:
                        continue
                    
                    # Merge cells if needed
                    if rowspan > 1 or colspan > 1:
                        word_table.cell(target_row, target_col).merge(
                            word_table.cell(target_row + rowspan - 1, target_col + colspan - 1)
                        )
                    
                    # Add content to cell
                    word_cell = word_table.cell(target_row, target_col)
                    word_cell.text = cell_text
                    
                    # Apply basic styling
                    if cell.name == 'th':
                        word_cell.paragraphs[0].runs[0].bold = True
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Move to next column
                    col_offset += colspan
                    
        except Exception as e:
            print(f"Error parsing HTML table: {e}")
            # Fallback: add as plain text
            doc.add_paragraph(str(element))
    
    def detect_structure(self, text):
        lines = text.split('\n')
        structured = []

        for line in lines:
            line = line.strip()

            if re.match(r'^(\d+\.|\-|\*)\s+', line):
                structured.append(('list', line))
            elif len(line) < 60 and line.isupper():
                structured.append(('heading', line))
            else:
                structured.append(('para', line))

        return structured
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # DOCX PIPELINE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _pipeline_docx(self, path: str, out: str) -> dict:
        source_doc = Document(path)
        new_doc    = Document()
        self._h4_counters = {}  

        # ── STEP 1: EXTRACT ──────────────────────────────────────────────
        
        raw_lines = []

        # ── STEP 1: EXTRACT (Preserving Order) ──────────────────────────
        for elem in source_doc.element.body.iterchildren():
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            
            if tag == 'p':
                # Walk ancestor chain looking for Fallback tag
                is_fallback = False
                parent = elem.getparent()
                while parent is not None:
                    if (parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag) == 'Fallback':
                        is_fallback = True; break
                    parent = parent.getparent()
                if is_fallback: continue

                para       = Paragraph(elem, source_doc)
                style_name = para.style.name if para.style else 'Normal'
                
                # Word list-numbering
                pPr      = elem.find(qn('w:pPr'))
                numPr    = pPr.find(qn('w:numPr'))   if pPr   is not None else None
                numId_el = numPr.find(qn('w:numId')) if numPr is not None else None
                ilvl_el  = numPr.find(qn('w:ilvl'))  if numPr is not None else None
                num_id   = int(numId_el.get(qn('w:val'), 0)) if numId_el is not None else 0
                num_lvl  = int(ilvl_el.get(qn('w:val'),  0)) if ilvl_el  is not None else 0

                # Run-level signals: bold and size
                run_sizes, is_bold = [], False
                for r in elem.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            try: run_sizes.append(int(sz.get(qn('w:val'))))
                            except: pass
                        if rPr.find(qn('w:b')) is not None: is_bold = True
                run_size = max(run_sizes) if run_sizes else 0

                text_full = para.text.strip()
                if text_full:
                    raw_lines.append({
                        'type': 'text', 'text': text_full, 'style': style_name,
                        'is_bold': is_bold, 'run_size': run_size,
                        'num_id': num_id, 'num_lvl': num_lvl
                    })

            elif tag == 'tbl':
                # Convert table to HTML for chat preview
                table_html = "<table>"
                for row in elem.findall(qn('w:tr')):
                    table_html += "<tr>"
                    for cell in row.findall(qn('w:tc')):
                        # Get all text from all paragraphs in the cell
                        cell_text = ""
                        for p in cell.findall(qn('w:p')):
                            for r in p.findall(qn('w:r')):
                                t = r.find(qn('w:t'))
                                if t is not None: cell_text += t.text
                            cell_text += " "
                        table_html += f"<td>{cell_text.strip()}</td>"
                    table_html += "</tr>"
                table_html += "</table>"

                raw_lines.append({
                    'type': 'table', 'node': elem, 'text': table_html
                })

        # ── STEP 2: JOIN ─────────────────────────────────────────────────
        # Word COM PDF→DOCX conversion fragments numbered items across paragraphs:

        joined = []
        i = 0
        while i < len(raw_lines):
            item = raw_lines[i]
            t    = item['text'].strip()

            # Lone digit — attempt forward stitch
            if re.match(r'^\d+$', t) and i + 1 < len(raw_lines):
                next_t = raw_lines[i + 1]['text'].strip()

                if re.match(r'^[\.\)]\s+\S', next_t):
                    # "1" + ". Customer opens..." → "1. Customer opens..."
                    # Normalize both "." and ")" separators to ". "
                    normalized = re.sub(r'^[\.\)]\s*', '. ', next_t)
                    item = dict(item)
                    item['text'] = t + normalized
                    joined.append(item)
                    i += 2
                    continue

                elif next_t in ('.', ')') and i + 2 < len(raw_lines):
                    # "3" + ")" + ". System validates..." → "3. System validates..."
                    rest       = raw_lines[i + 2]['text'].strip()
                    rest_clean = re.sub(r'^[\.\)]\s*', '', rest)
                    item = dict(item)
                    item['text'] = t + '. ' + rest_clean
                    joined.append(item)
                    i += 3
                    continue

                elif re.match(r'^\d+$', next_t):
                    # Two consecutive lone digits — skip this one
                    i += 1
                    continue

                else:
                    # Lone digit with no joinable continuation — keep as placeholder
                    item = dict(item)
                    item['text'] = t + '.'
                    joined.append(item)
                    i += 1
                    continue

            # Lone punctuation — append to previous only if it isn't already a complete numbered item
            elif joined and len(t) == 1 and t in '.)':
                if not re.match(r'^\d+[\.\)]', joined[-1]['text']):
                    joined[-1]['text'] += t
                i += 1
                continue

            # Fragment like ". Customer scans QR on table"
            elif joined and re.match(r'^[\.\)]\s+\S', t):
                prev = joined[-1]['text']
                if not re.match(r'^\d+[\.\)]\s+\S', prev):
                    # Previous is not yet a complete numbered item — stitch
                    joined[-1]['text'] += t
                    i += 1
                    continue
                else:
                    # Previous is already a complete numbered item — this is a
                    # continuation bullet from the same step; strip leading dot
                    item = dict(item)
                    item['text'] = re.sub(r'^[\.\)]\s*', '', t).strip()
                    if item['text']:
                        joined.append(item)
                    i += 1
                    continue

            # Lone "-" paragraph — prefix the NEXT line as a dash bullet
            elif t == '-':
                if i + 1 < len(raw_lines):
                    raw_lines[i + 1]['text'] = '- ' + raw_lines[i + 1]['text'].lstrip('- ')
                i += 1
                continue

            if t:
                joined.append(item)
            i += 1

        raw_lines = joined

        # ── STEPS 3 & 4: CLASSIFY + BUILD ──────────────────────────────────
        self._build_from_signals(raw_lines, new_doc)

        branded_doc = self.style_manager.apply_template_styles(new_doc)
        self._apply_final_features(branded_doc)
        
        branded_doc.save(out)
        return {'success': True, 'paragraphs': len(new_doc.paragraphs)}

    def _build_from_signals(self, raw_lines, doc):
        """Standard processing logic shared by initial files and text edits."""
        expecting_list = False
        is_first_non_empty = True

        for item in raw_lines:
            if item.get('type') == 'table':
                # RESTORE TABLE PROCESSING
                from docx.table import Table
                new_tbl = Table(item['node'], doc)
                doc.element.body.append(new_tbl._element)
                continue

            text     = item['text'].replace('Ł', '').strip()
            style    = item.get('style', 'Normal')
            is_bold  = item.get('is_bold', False)
            run_size = item.get('run_size', 0)
            num_id   = item.get('num_id', 0)
            num_lvl  = item.get('num_lvl', 0)
            is_numbered_list_item = num_id > 0

            if not text: continue
            
            # Signals
            is_section_label    = text.endswith(':') and len(text.split()) <= 6
            is_numbered_heading = bool(re.match(r'^\d+[\.\)]\s+[A-Z]', text))
            is_explicit_bullet  = text.startswith(('☐', '☑', '•', '▪', '►', '■', '□'))
            is_dash_bullet      = bool(re.match(r'^[-–]\s+\S', text))

            # --- Title Promotion ---
            if is_first_non_empty and len(text.split()) <= 15 and not text.endswith('.'):
                doc.add_heading(text, level=1)
                is_first_non_empty = False
                continue

            is_first_non_empty = False

            # --- Classifier logic (H1-H4) ---
            if 'Heading 1' in style: doc.add_heading(text, level=1)
            elif 'Heading 2' in style or is_numbered_heading or run_size >= 26 or (is_bold and run_size >= 22):
                doc.add_heading(text, level=2)
            elif 'Heading 3' in style or is_section_label or (is_bold and len(text.split()) <= 8):
                doc.add_heading(text, level=3)
            elif is_numbered_list_item and num_lvl == 0:
                self._h4_counters[num_id] = self._h4_counters.get(num_id, 0) + 1
                doc.add_heading(f"{self._h4_counters[num_id]}. {text}", level=4)
            elif is_explicit_bullet or is_dash_bullet or expecting_list:
                clean = re.sub(r'^[-–☐☑•▪►■□]\s*', '', text).strip().rstrip('-').strip()
                if clean:
                    try: doc.add_paragraph(clean, style='List Bullet')
                    except: doc.add_paragraph(clean)
            else:
                doc.add_paragraph(text)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PDF PIPELINE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _pipeline_pdf(self, path: str, out: str) -> dict:
        """Adobe Extract API → semantic JSON elements → branded DOCX."""
        try:
            elements    = adobe_pdf_extract(path, ADOBE_CLIENT_ID, ADOBE_CLIENT_SECRET)
            doc         = self._build_from_adobe_json(elements)
            branded_doc = self.style_manager.apply_template_styles(doc)
            
            # Apply optional features
            self._apply_final_features(branded_doc)

            branded_doc.save(out)
            return {'success': True}
        except Exception as e:
            print(f'[DocumentProcessor] Adobe API error: {e}')
            return {'success': False, 'error': str(e)}

    def _build_from_adobe_json(self, elements: list) -> Document:
        """Convert Adobe Extract API semantic elements into a structured DOCX."""
        doc            = Document(self.template_path)
        expecting_list = False

        for el in elements:
            path = el.get('Path', '')
            text = el.get('Text', '').strip()

            if not text and 'Table' not in path:
                continue

            # Adobe separates bullet labels ("/Lbl") from content ("/LBody").
            # Skip labels — LBody carries the full text we need.
            if '/Lbl' in path:
                continue

            text = text.replace('Ł', '').strip()

            is_section_label = text.endswith(':') and len(text.split()) <= 6
            is_explicit_bullet = text.startswith(('-', '☐', '•', '▪', '➤', '■'))

            # Update list context
            if is_section_label:
                expecting_list = True
            elif text.endswith('.') or len(text) > 120 or re.search(r'/H\d', path):
                expecting_list = False

            if '/Title' in path:
                doc.add_heading(text, level=1)
            elif re.search(r'/H1', path):
                doc.add_heading(text, level=1)
            elif re.search(r'/H2', path):
                doc.add_heading(text, level=2)
            elif re.search(r'/H3', path) or is_section_label:
                doc.add_heading(text, level=3)
            elif '/LI' in path or '/LBody' in path or is_explicit_bullet or expecting_list:
                clean = re.sub(r'^[-\[\]☐•▪➤■]\s*', '', text)
                doc.add_paragraph(clean, style='List Bullet')
            elif 'Table' in path:
                pass  # table reconstruction not yet implemented
            else:
                doc.add_paragraph(text)

        return doc


