"""
CoverPageManager — Injects a cover page at the start of the document.

INJECTION STRATEGY:
  Elements are inserted at index 0 in reverse order (stack behaviour).
  Insert bottom-most element first so the final page order is:
    Title → Subtitle → Executive Summary → Date → Section break → Content
"""

import ollama
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from copy import deepcopy


class CoverPageManager:

    def __init__(self, model_name: str = 'phi3'):
        self.model_name = model_name
        self.bullet_prefixes = ('•', '-', '*', '▪', '►', '➤', '→', '○', '●', '‣', '–', '—')

    def create_cover_page(self, doc: Document) -> bool:
        """
        Main entry point. Extracts title, generates AI summary, and injects
        the cover page at the very top of the document.
        """
        title, subtitle = self._extract_title_subtitle(doc)
        if not title:
            print('[Cover Page] No title found, skipping')
            return False

        print(f'[Cover Page] Title: {title[:50]}')
        doc_text = self._extract_document_text(doc)
        summary  = self._generate_summary_with_ollama(doc_text, title)

        body = doc.element.body

        # Insert in reverse order — each insert(0) pushes previous down.
        # Final page order: Title → Subtitle → Summary → Date → section break → content

        # 1. Section break first (becomes last on page, pushes content to page 2)
        body.insert(0, self._make_cover_section_break(doc))

        # 2. Date
        body.insert(0, self._create_centered_para(doc, datetime.now().strftime('%B %d, %Y'), 11)._element)

        # 3. Executive Summary
        if summary:
            body.insert(0, self._create_centered_para(doc, summary, 11, italic=True)._element)
            body.insert(0, self._create_centered_para(doc, 'Executive Summary', 13, bold=True)._element)

        # 4. Subtitle
        if subtitle:
            body.insert(0, self._create_centered_para(doc, subtitle, 14)._element)

        # 5. Title (ends up at index 0 — very top)
        body.insert(0, self._create_centered_para(doc, title.upper(), 22, bold=True)._element)

        print('[Cover Page] ✅ Injection complete')
        return True

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # XML HELPERS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _create_centered_para(self, doc: Document, text: str, size: int,
                               bold: bool = False, italic: bool = False):
        """
        Builds a centered paragraph, extracts its XML element, then removes it
        from the document body. Only the element is kept for direct insertion.
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._force_para_center_xml(p)
        p.paragraph_format.space_after = Pt(12)

        r = p.add_run(text)
        r.font.size      = Pt(size)
        r.font.bold      = bold
        r.font.italic    = italic
        r.font.name      = 'Calibri'
        r.font.color.rgb = RGBColor(0, 0, 0)

        # Remove from doc body — we only want the XML element for manual insertion
        p._element.getparent().remove(p._element)
        return p

    def _make_cover_section_break(self, doc: Document):
        """
        Paragraph containing a next-page section break with vertical centering.
        Copies header/footer references from the template so the cover page
        gets the company letterhead too.
        """
        p      = OxmlElement('w:p')
        pPr    = OxmlElement('w:pPr')
        sectPr = OxmlElement('w:sectPr')

        # Copy header/footer references and page size/margins from template
        template_sectPr = doc.sections[0]._sectPr
        for tag in ('w:headerReference', 'w:footerReference', 'w:pgSz', 'w:pgMar'):
            for elem in template_sectPr.findall(qn(tag)):
                sectPr.append(deepcopy(elem))

        # Vertically center the cover page content
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        sectPr.append(vAlign)

        # Next page break so content starts on page 2
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'nextPage')
        sectPr.append(sectType)

        pPr.append(sectPr)
        p.append(pPr)
        return p

    def _force_para_center_xml(self, para):
        """Write center alignment directly to XML so it cannot be overridden by styles."""
        pPr = para._element.get_or_add_pPr()
        for existing in pPr.findall(qn('w:jc')):
            pPr.remove(existing)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # AI SUMMARY + TEXT EXTRACTION
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _generate_summary_with_ollama(self, doc_text: str, title: str) -> str:
        """Generate a 1-2 sentence executive summary using the local Ollama model."""
        if len(doc_text) > 6000:
            doc_text = doc_text[:6000]
        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': f'Summarize in 1-2 sentences:\nTitle: {title}\n\n{doc_text}'}],
                options={'temperature': 0.3, 'num_predict': 150}
            )
            return response['message']['content'].strip()
        except Exception:
            return 'Technical documentation and strategic project overview.'

    def _extract_title_subtitle(self, doc: Document) -> tuple:
        """
        Scan the first 15 paragraphs for the title and subtitle.
        Skips empty lines, TOC headings, and bullet points.
        Returns (title, subtitle) — either can be None.
        """
        candidates = []
        for para in doc.paragraphs[:15]:
            text = para.text.strip()
            if not text:
                continue
            if text.upper() == 'TABLE OF CONTENTS':
                continue
            if any(text.startswith(prefix) for prefix in self.bullet_prefixes):
                continue
            if 2 <= len(text.split()) <= 25:
                candidates.append(text)
            if len(candidates) >= 2:
                break
        return (
            candidates[0] if candidates else None,
            candidates[1] if len(candidates) > 1 else None
        )

    def _extract_document_text(self, doc: Document) -> str:
        """Return all non-empty paragraph text joined by newlines."""
        return '\n'.join(p.text.strip() for p in doc.paragraphs if p.text.strip())