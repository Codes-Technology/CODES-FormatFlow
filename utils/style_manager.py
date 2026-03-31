"""
StyleManager 
"""

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from copy import deepcopy


# Relationship namespace
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


class StyleManager:
    def __init__(self, template_path: str, font_family: str = 'Calibri', font_size: int = 11):
        self.template_path = template_path
        self.template = Document(template_path)
        
        # User-selected settings
        self.font_family = font_family
        self.BODY_FONT = font_family 
        self.font_size = font_size
        
        self.BODY_COLOR = "000000"
        # Calculate sizes based on user's choice (in half-points)
        self.BODY_SIZE = font_size * 2              # 11pt → 22
        self.MIN_SIZE = self.BODY_SIZE
        self.H1_SIZE = int(font_size * 1.45 * 2)    # 11pt → 16pt (32 half-pts)
        self.H2_SIZE = int(font_size * 1.27 * 2)    # 11pt → 14pt (28 half-pts)
        self.H3_SIZE = int(font_size * 1.18 * 2)    # 11pt → 13pt (26 half-pts)
        self.H4_SIZE = int(font_size * 1.09 * 2)

    def apply_template_styles(self, doc: Document) -> Document:
        """Single entry point — full branding pipeline."""
        print('[StyleManager] Applying branding...')
        self._remove_document_protection(doc)
        self._remove_leading_empty_paragraphs(doc)
        self._override_style_sizes(doc)
        self._brand_body_paragraphs(doc)
        self._apply_page_layout_rules(doc)
        self._apply_heading_indents(doc)
        # self._add_page_number(doc)   # must run BEFORE _copy_header_footer
        # self._copy_header_footer(doc)
       

        print('[StyleManager] ✓ Done')
        return doc

    
    # Style-level sizes (styles.xml, not run-by-run)
    

    def _override_style_sizes(self, doc: Document):
        config = {
            'Normal':       {'sz': self.BODY_SIZE, 'bold': False},
            # Paragraph heading styles
            'Heading1':     {'sz': self.H1_SIZE,   'bold': True},
            'Heading 1':    {'sz': self.H1_SIZE,   'bold': True},
            'Heading2':     {'sz': self.H2_SIZE,   'bold': True},
            'Heading 2':    {'sz': self.H2_SIZE,   'bold': True},
            'Heading3':     {'sz': self.H3_SIZE,   'bold': True},
            'Heading 3':    {'sz': self.H3_SIZE,   'bold': True},
            # Linked CHARACTER styles
            'Heading1Char': {'sz': self.H1_SIZE,   'bold': True},
            'Heading2Char': {'sz': self.H2_SIZE,   'bold': True},
            'Heading3Char': {'sz': self.H3_SIZE,   'bold': True},
            'Heading4':     {'sz': self.H4_SIZE,   'bold': True},
            'Heading 4':    {'sz': self.H4_SIZE,   'bold': True},
            'Heading4Char': {'sz': self.H4_SIZE,   'bold': True},
        }

        styles_elem = doc.part.styles._element
        for style_el in styles_elem.findall(qn('w:style')):
            style_id   = style_el.get(qn('w:styleId'), '')
            name_el    = style_el.find(qn('w:name'))
            style_name = name_el.get(qn('w:val'), '') if name_el is not None else ''

            cfg = config.get(style_id) or config.get(style_name)
            if cfg is None:
                continue

            rPr = style_el.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style_el.append(rPr)

            for t in ('w:sz', 'w:szCs'):
                for e in rPr.findall(qn(t)):
                    rPr.remove(e)
            sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(cfg['sz']));   rPr.append(sz)
            szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(cfg['sz'])); rPr.append(szCs)

            for e in rPr.findall(qn('w:color')): rPr.remove(e)
            c = OxmlElement('w:color'); c.set(qn('w:val'), '000000'); rPr.append(c)

            for e in rPr.findall(qn('w:rFonts')): rPr.remove(e)
            f = OxmlElement('w:rFonts')
            f.set(qn('w:ascii'), self.font_family); f.set(qn('w:hAnsi'), self.font_family)
            f.set(qn('w:cs'), self.font_family);   rPr.append(f)

            for e in rPr.findall(qn('w:b')): rPr.remove(e)
            if cfg['bold']:
                rPr.append(OxmlElement('w:b'))

            print(f'  [StyleManager] Style "{style_id or style_name}" → '
                f'{cfg["sz"] / 2}pt {self.font_family} bold={cfg["bold"]}')
            
            pPr = style_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style_el.append(pPr)

            for e in pPr.findall(qn('w:spacing')):
                pPr.remove(e)

            spacing = OxmlElement('w:spacing')

            if cfg['bold']:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '120')
                spacing.set(qn('w:line'), '240')
                spacing.set(qn('w:lineRule'), 'auto')
            else:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '20')
                spacing.set(qn('w:line'), '220')
                spacing.set(qn('w:lineRule'), 'auto')

            pPr.append(spacing)

            if cfg['bold']:   
                for e in pPr.findall(qn('w:keepNext')):
                    pPr.remove(e)
                pPr.append(OxmlElement('w:keepNext'))

    def _apply_heading_indents(self, doc):
        """Apply indentation to headings and body text using fixed baselines"""
        
        current_baseline = 0
        
        for para in doc.paragraphs:
            style = para.style.name
            pPr = para._p.get_or_add_pPr()
            
            # Remove old indent
            for e in pPr.findall(qn('w:ind')):
                pPr.remove(e)
            
            # 1. Update baseline based on headings
            if 'Heading 1' in style or style == 'Heading1':
                indent_left = 0
                current_baseline = 360 # Body text under H1 is at 0.25"
            elif 'Heading 2' in style or style == 'Heading2':
                indent_left = 360 # 0.25 inch
                current_baseline = 720 # Body text under H2 is at 0.5"
            elif 'Heading 3' in style or style == 'Heading3':
                indent_left = 720 # 0.5 inch
                current_baseline = 1080 # Body text under H3 is at 0.75"
            elif 'Heading 4' in style or style == 'Heading4':
                indent_left = 960 # 0.67 inch
                current_baseline = 1200 # Body text under H4 is at 0.83"
            
            # 2. Apply baseline for non-headings
            else:
                indent_left = current_baseline
            
            # 3. Handle Lists (fixed offset from baseline, no accumulation)
            if 'List' in style:
                # Use current_baseline (e.g. 1080 for H3) as the reference point
                # Bullet itself stays at the baseline (where text normally starts)
                # But we use hanging indent to pull the bullet character left
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), str(indent_left + 360)) # Increase overall left margin
                ind.set(qn('w:hanging'), '280') # Pull bullet back 
                pPr.append(ind)
                
                # Add explicit tab stop so bullet text aligns properly
                for e in pPr.findall(qn('w:tabs')):
                    pPr.remove(e)
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'left')
                tab.set(qn('w:pos'), str(indent_left + 360))
                tabs.append(tab)
                pPr.append(tabs)
            
            else:
                # Normal paragraph indent
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), str(indent_left))
                pPr.append(ind)
            
        # Push all Tables slightly inward to match body text flow
        for table in doc.tables:
            tbl_pr = table._element.tblPr
            tblInd = tbl_pr.xpath('w:tblInd')
            if not tblInd:
                tblInd = OxmlElement('w:tblInd')
                tblInd.set(qn('w:w'), '360')
                tblInd.set(qn('w:type'), 'dxa')
                tbl_pr.append(tblInd)
            else:
                tblInd[0].set(qn('w:w'), '360')

    def _copy_header_footer(self, doc: Document):
        template_section = self.template.sections[0]
        
        # Reload template fresh each time to avoid stale relationships
        self.template = Document(self.template_path)
        template_section = self.template.sections[0]
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            self._deep_copy_hdr_ftr(template_section.header, section.header)
            self._deep_copy_hdr_ftr(template_section.footer, section.footer)
        # self._fix_duplicate_media(doc)
            
    def _add_page_number(self, doc: Document):
        
        for section in doc.sections:
            footer = section.footer

            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = p.add_run()

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            instrText = OxmlElement('w:instrText')
            instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            instrText.text = " PAGE "
            run._r.append(instrText)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)

    def _deep_copy_hdr_ftr(self, source, target):
        rId_map = {}
        for old_rId, rel in source.part.rels.items():
            try:
                if 'image' in rel.reltype:
                    new_rId = target.part.relate_to(rel.target_part, rel.reltype)
                    rId_map[old_rId] = new_rId
                elif 'hyperlink' in rel.reltype:
                    new_rId = target.part.relate_to(
                        rel.target_ref, rel.reltype, is_external=True)
                    rId_map[old_rId] = new_rId
            except Exception as e:
                print(f'  [StyleManager] Rel copy warning ({old_rId}): {e}')

        target_elem = target._element
        for p in list(target_elem.findall(qn('w:p'))):
            target_elem.remove(p)

        for src_p in source._element.findall(qn('w:p')):
            new_p = deepcopy(src_p)
            if rId_map:
                _remap_rids(new_p, rId_map)
            target_elem.append(new_p)
    
    def _iter_all_paragraphs(self, doc: Document):
        """
        Yield every paragraph in the document,
        including paragraphs inside tables.
        """
        # Normal paragraphs
        for para in doc.paragraphs:
            yield para

        # Table paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    def _brand_body_paragraphs(self, doc: Document):
        heading_names = {
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
            'Heading1',  'Heading2',  'Heading3',  'Heading4',
        }

        for para in self._iter_all_paragraphs(doc):
            style_name = para.style.name if para.style else 'Normal'
            is_heading = style_name in heading_names or style_name.startswith('Heading')

            for r in para._p.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)

                # ── Color: always black ──
                for e in rPr.findall(qn('w:color')):
                    rPr.remove(e)
                ce = OxmlElement('w:color')
                ce.set(qn('w:val'), self.BODY_COLOR)
                rPr.append(ce)

                # ── Font family: always apply ──
                for e in rPr.findall(qn('w:rFonts')):
                    rPr.remove(e)
                fe = OxmlElement('w:rFonts')
                fe.set(qn('w:ascii'), self.BODY_FONT)
                fe.set(qn('w:hAnsi'), self.BODY_FONT)
                rPr.append(fe)

                # ── Remove italic everywhere ──
                for e in rPr.findall(qn('w:i')):
                    rPr.remove(e)

                if is_heading:
                    # HEADINGS: remove run-level size overrides
                    # so the STYLE-level sizes (H1=15.5pt etc.) take effect
                    for e in rPr.findall(qn('w:sz')):
                        rPr.remove(e)
                    for e in rPr.findall(qn('w:szCs')):
                        rPr.remove(e)
                else:
                    # BODY TEXT: force to BODY_SIZE, remove bold
                    for e in rPr.findall(qn('w:b')):
                        rPr.remove(e)
                    for e in rPr.findall(qn('w:sz')):
                        rPr.remove(e)
                    for e in rPr.findall(qn('w:szCs')):
                        rPr.remove(e)
                    s = OxmlElement('w:sz')
                    s.set(qn('w:val'), str(self.BODY_SIZE))
                    rPr.append(s)
                    sc = OxmlElement('w:szCs')
                    sc.set(qn('w:val'), str(self.BODY_SIZE))
                    rPr.append(sc)

    def _apply_page_layout_rules(self, doc: Document):
        first_h1 = True
        prev_style = None

        for para in doc.paragraphs:
            if not para.style:
                continue

            style_name = para.style.name

            if style_name in ('Heading 1', 'Heading1'):
                if not first_h1:
                    
                    if prev_style and not prev_style.startswith('Heading'):
                        para.paragraph_format.page_break_before = True
                first_h1 = False

            if style_name.startswith('Heading'):
                para.paragraph_format.keep_with_next = True
                pPr = para._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    para._p.insert(0, pPr)
                if pPr.find(qn('w:keepNext')) is None:
                    pPr.append(OxmlElement('w:keepNext'))

            if para.text.strip():
                prev_style = style_name
    
    def _remove_leading_empty_paragraphs(self, doc):
        while doc.paragraphs and doc.paragraphs[0].text.strip() == "":
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

    def _remove_document_protection(self, doc: Document):
        """Remove all document protection to make it editable"""
        print("  Removing document protection...")
        
        try:
            # Remove document protection from settings
            settings = doc.settings.element
            
            # Remove documentProtection element
            for protect in settings.findall(qn('w:documentProtection')):
                settings.remove(protect)
            
            # Remove writeProtection element (read-only)
            for write_protect in settings.findall(qn('w:writeProtection')):
                settings.remove(write_protect)
            
            # Unlock all content controls
            for content_control in doc.element.xpath('.//w:sdt'):
                sdtPr = content_control.find(qn('w:sdtPr'))
                if sdtPr is not None:
                    for lock in sdtPr.findall(qn('w:lock')):
                        sdtPr.remove(lock)
            
            print("    ✓ Document is now editable")
        except Exception as e:
            print(f"    ⚠ Error: {e}")
    

def _remap_rids(element, rId_map: dict):
    
    for attr_key in list(element.attrib.keys()):
        local = attr_key.split('}')[-1] if '}' in attr_key else attr_key
        ns    = attr_key.split('}')[0].lstrip('{') if '}' in attr_key else ''

        if ns == R_NS and local in ('embed', 'id', 'link', 'href', 'pict'):
            old_val = element.attrib[attr_key]
            if old_val in rId_map:
                element.attrib[attr_key] = rId_map[old_val]

    for child in element:
        _remap_rids(child, rId_map)