import os
import datetime
from typing import List, Dict
from docx import Document

class BatchProcessor:
    """Process multiple files in batch and merge them cleanly."""
    
    def __init__(self, output_folder: str):
        self.output_folder = output_folder

    def process_and_merge(self, file_list: List[str], processor) -> Dict:
        """Merges multiple files into ONE single document using docxcompose."""
        print("\n" + "="*70)
        print(f"BATCH MERGE: {len(file_list)} FILES")
        print("="*70)
        
        try:
            processed_temp_files = []
            successful_count = 0
            
            # Step 1: Extract RAW content individually
            for i, file_path in enumerate(file_list):
                filename = os.path.basename(file_path)
                print(f"[{i+1}/{len(file_list)}] Processing: {filename}")
                
                temp_inter = os.path.join(self.output_folder, f"temp_{datetime.datetime.now().timestamp()}_{i}.docx")
                result = processor.universal_extract(file_path, temp_inter)
                
                if result['success']:
                    processed_temp_files.append(temp_inter)
                    successful_count += 1
                    
                    # Apply table reconstruction to intermediary file before merging
                    ext = os.path.splitext(file_path)[1].lower()
                    if ext in ['.pdf', '.txt']:
                        try:
                            temp_doc = Document(temp_inter)
                            processor.style_manager.reconstruct_tables(temp_doc)
                            temp_doc.save(temp_inter)
                            print(f"  [Style] Applied table reconstruction for {ext}")
                        except Exception as e:
                            print(f"  [Warning] Table reconstruction failed for {filename}: {e}")
                else:
                    print(f"  [X] Extraction failed: {result.get('error')}")

            if not processed_temp_files:
                return {'success': False, 'error': 'No valid content extracted from files'}

            # Step 2: Prepare Master Document (Safe Clear)
            print("\nPreparing master document...")
            master_doc = Document(processor.template_path)
            
            for table in list(master_doc.tables):
                t = table._element
                if t.getparent() is not None:
                    t.getparent().remove(t)
            
            if len(master_doc.paragraphs) > 1:
                all_paras = list(master_doc.paragraphs)
                last_para = all_paras[-1]
                for para in all_paras:
                    if para is not last_para:
                        p = para._element
                        if p.getparent() is not None:
                            p.getparent().remove(p)
            
            if master_doc.paragraphs:
                master_doc.paragraphs[-1].text = ""

            # Step 3: Merge documents into Master
            print("Merging content into master...")
            from utils.merge_fix import safe_merge
            
            sub_docs = []
            for i, sub_doc_path in enumerate(processed_temp_files):
                sub_doc = Document(sub_doc_path)
                
                # Add page break rule for subsequent documents
                if i > 0 and sub_doc.paragraphs:
                    first_real_para = next((p for p in sub_doc.paragraphs if p.text.strip()), sub_doc.paragraphs[0])
                    first_real_para.paragraph_format.page_break_before = True
                        
                sub_docs.append(sub_doc)

            final_doc = safe_merge(master_doc, sub_docs)
            print("  [OK] Merged content successfully")
            
            # Step 4: Final Styles
            print("\nApplying final styles and branding...")
            final_doc = processor.style_manager.apply_template_styles(final_doc)
            
            # Save Output
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_name = f"Merged_Document_{timestamp}.docx"
            output_path = os.path.join(self.output_folder, output_name)
            final_doc.save(output_path)
            
            print(f"  [OK] Created: {output_name}")
            self.cleanup_temp_files(processed_temp_files)
            
            return {'success': True, 'output_file': output_name, 'count': successful_count}
            
        except ImportError:
            return {'success': False, 'error': 'docxcompose not installed. Run: pip install docxcompose'}
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def cleanup_temp_files(self, file_list: List[str]):
        """Remove temporary files."""
        for file_path in file_list:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception:
                pass